"""
Resume Parser Module - Extract data from PDF/DOCX files
Handles accurate data extraction without hallucination
"""

import re
import json
from typing import Dict, List, Optional
from pathlib import Path
import PyPDF2
from docx import Document


class ResumeParser:
    """Parse resume files and extract structured data"""
    
    # Common technical skills database
    TECHNICAL_SKILLS = {
        "Programming Languages": [
            "Python", "JavaScript", "Java", "C++", "C#", "PHP", "Ruby", "Go", 
            "Rust", "Swift", "Kotlin", "TypeScript", "Scala", "Perl", "R",
            "MATLAB", "Objective-C", "Groovy", "Haskell", "Elixir"
        ],
        "Web Frameworks": [
            "React", "Vue.js", "Angular", "Django", "Flask", "Node.js", "Express",
            "Spring", "ASP.NET", "Laravel", "FastAPI", "Fastify", "NestJS",
            "Next.js", "Nuxt.js", "Ruby on Rails"
        ],
        "Databases": [
            "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Cassandra",
            "DynamoDB", "Firebase", "Oracle", "SQLite", "Elasticsearch",
            "Neo4j", "CouchDB", "MariaDB"
        ],
        "Cloud Platforms": [
            "AWS", "Azure", "Google Cloud", "Heroku", "Digital Ocean",
            "Linode", "CloudFlare", "IBM Cloud", "Oracle Cloud", "Alibaba Cloud"
        ],
        "Tools & Platforms": [
            "Git", "Docker", "Kubernetes", "Jenkins", "GitHub", "GitLab",
            "Bitbucket", "JIRA", "Slack", "Linux", "Windows", "MacOS",
            "Nginx", "Apache", "Terraform", "Ansible", "Docker Compose"
        ],
        "Data & ML": [
            "Machine Learning", "Tensorflow", "PyTorch", "Scikit-learn",
            "Pandas", "NumPy", "Data Analysis", "Statistical Analysis",
            "Deep Learning", "Natural Language Processing", "Computer Vision",
            "Apache Spark", "Hadoop", "Big Data"
        ],
        "Soft Skills": [
            "Communication", "Leadership", "Problem Solving", "Team Work",
            "Project Management", "Agile", "Scrum", "Time Management",
            "Critical Thinking", "Collaboration", "Presentation", "Documentation"
        ],
        "Testing & QA": [
            "Unit Testing", "Integration Testing", "Testing", "QA",
            "Selenium", "Jest", "Pytest", "Mocha", "Chai", "JUnit"
        ],
        "DevOps": [
            "CI/CD", "System Design", "Microservices", "Scalability",
            "High Availability", "Performance Optimization", "Security"
        ]
    }
    
    # Flatten skills for quick lookup
    ALL_SKILLS = []
    
    def __init__(self):
        """Initialize parser with skill database"""
        for category_skills in self.TECHNICAL_SKILLS.values():
            self.ALL_SKILLS.extend(category_skills)
        self.ALL_SKILLS = list(set([s.lower() for s in self.ALL_SKILLS]))
    
    def parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    def parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    def parse_resume(self, file_path: str) -> str:
        """Parse resume file (PDF or DOCX) and return text"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            return self.parse_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            return self.parse_docx(str(file_path))
        else:
            return "Unsupported file format"
    
    def get_page_count(self, file_path: str) -> int:
        """Get actual page count from resume"""
        try:
            file_path = Path(file_path)
            
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    return len(pdf_reader.pages)
            elif file_path.suffix.lower() == '.docx':
                # For DOCX, estimate based on content length
                # DOCX doesn't have true page concept like PDF
                doc = Document(file_path)
                text_length = sum(len(p.text) for p in doc.paragraphs)
                # Rough estimate: ~3000 characters per page
                estimated_pages = max(1, (text_length + 2999) // 3000)
                return estimated_pages
            return 0
        except Exception as e:
            return 0
    
    def extract_name(self, text: str) -> str:
        """Extract full name from resume"""
        # Common patterns for name at the beginning
        lines = text.strip().split('\n')
        
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line.split()) >= 2:
                # Check if it looks like a name (not an email or phone)
                if '@' not in line and '+' not in line and '(' not in line:
                    words = line.split()
                    if len(words) >= 2:
                        return line
        
        return "Not Found"
    
    def extract_email(self, text: str) -> str:
        """Extract email address from resume"""
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else "Not Found"
    
    def extract_phone(self, text: str) -> str:
        """Extract phone number from resume"""
        # Multiple phone patterns
        phone_patterns = [
            r'\+?1?\s?\(?[\d\s\-\(\)]{9,}\d',  # US/International format
            r'[\d\s\-\(\)]{9,}',  # General format
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text[:500])  # Check first 500 chars
            if matches:
                return matches[0].strip()
        
        return "Not Found"
    
    def extract_education(self, text: str) -> List[Dict]:
        """Extract education details from resume"""
        education = []
        
        # Education keywords
        education_keywords = [
            'education', 'degree', 'bachelor', 'master', 'phd', 'diploma',
            'graduate', 'undergraduate', 'university', 'college', 'institute',
            'b.tech', 'b.s', 'm.s', 'b.a', 'm.a', 'b.e', 'm.tech'
        ]
        
        # Find education section
        text_lower = text.lower()
        education_idx = -1
        
        for keyword in education_keywords:
            idx = text_lower.find(keyword)
            if idx != -1 and idx < 5000:  # Usually near start
                education_idx = max(education_idx, idx)
        
        if education_idx == -1:
            return education
        
        # Extract education section
        section = text[education_idx:education_idx+2000]
        
        # Look for degree patterns
        degree_patterns = [
            r'(Bachelor|Master|PhD|B\.Tech|M\.Tech|B\.S|M\.S|B\.A|M\.A)\s+(?:of\s+)?(?:Science|Arts|Engineering|Technology|Commerce|Administration)?',
            r'(Associate|Diploma|Certificate)\s+(?:in\s+)?([^,\n]*)',
        ]
        
        for pattern in degree_patterns:
            matches = re.finditer(pattern, section, re.IGNORECASE)
            for match in matches:
                degree_text = match.group(0).strip()
                
                # Extract branch/field
                branch = "Not Mentioned"
                branch_patterns = [
                    r'(?:of|in)\s+([A-Za-z\s]+?)(?:\s+from|\s+at|,|$)',
                    r'in\s+([A-Za-z\s&]+)',
                ]
                
                for bp in branch_patterns:
                    b_match = re.search(bp, degree_text, re.IGNORECASE)
                    if b_match:
                        branch = b_match.group(1).strip()
                        break
                
                # Extract university
                university = "Not Mentioned"
                univ_patterns = [
                    r'from\s+([A-Za-z\s,\.]+?)(?:,|$)',
                    r'at\s+([A-Za-z\s,\.]+?)(?:,|$)',
                ]
                
                # Look ahead for university
                end_idx = section.find(match.group(0)) + len(match.group(0))
                lookahead = section[end_idx:end_idx+200]
                
                for up in univ_patterns:
                    u_match = re.search(up, lookahead, re.IGNORECASE)
                    if u_match:
                        university = u_match.group(1).strip()
                        break
                
                if degree_text not in [e.get('degree') for e in education]:
                    education.append({
                        "degree": degree_text,
                        "branch": branch,
                        "university": university
                    })
        
        return education
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume accurately"""
        skills_found = []
        text_lower = text.lower()
        
        # Find skills section
        skills_keywords = ['skills', 'technical skills', 'competencies', 'expertise']
        skills_section_start = -1
        
        for keyword in skills_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                skills_section_start = idx
                break
        
        # If skills section found, prioritize it
        if skills_section_start != -1:
            # Extract text after skills section (until next section)
            section_end = min(skills_section_start + 2000, len(text))
            next_section_patterns = ['experience', 'education', 'projects', 'certification']
            
            for pattern in next_section_patterns:
                idx = text_lower.find(pattern, skills_section_start + 20)
                if idx != -1 and idx < section_end:
                    section_end = idx
            
            search_text = text[skills_section_start:section_end]
        else:
            search_text = text
        
        # Search for skills (case-insensitive)
        search_text_lower = search_text.lower()
        
        for skill in self.ALL_SKILLS:
            # Use word boundaries for accuracy
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, search_text_lower, re.IGNORECASE):
                # Find original case in full text
                original_match = re.search(pattern, text, re.IGNORECASE)
                if original_match:
                    # Get the properly cased skill name
                    matched_text = original_match.group(0)
                    # Find matching skill with original case
                    for full_skill in self.ALL_SKILLS:
                        if full_skill == skill:
                            # Try to get better casing from text
                            for orig_skill_list in self.TECHNICAL_SKILLS.values():
                                for orig_skill in orig_skill_list:
                                    if orig_skill.lower() == skill:
                                        if orig_skill not in skills_found:
                                            skills_found.append(orig_skill)
                                        break
                            break
        
        # Remove duplicates while preserving order
        skills_found = list(dict.fromkeys(skills_found))
        
        return skills_found if skills_found else []
    
    def extract_experience(self, text: str) -> List[Dict]:
        """Extract experience/work history from resume"""
        experience = []
        
        # Look for experience section
        text_lower = text.lower()
        exp_keywords = ['experience', 'work experience', 'professional experience', 'employment']
        
        exp_idx = -1
        for keyword in exp_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                exp_idx = idx
                break
        
        if exp_idx == -1:
            return experience
        
        # Get experience section
        section = text[exp_idx:exp_idx+3000]
        
        # Look for job entries (company + position pattern)
        # Usually formatted as Company/Position - Dates
        job_patterns = [
            r'([A-Za-z\s&\.]+)(?:\s*[\|\-\,]\s*)?([A-Za-z\s]+)?\s*(?:\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)',
        ]
        
        # Simple approach: Look for lines with company names
        lines = section.split('\n')
        
        for line in lines[:10]:  # Limit to avoid noise
            line = line.strip()
            if line and len(line) > 10 and not line.startswith('-'):
                # Check if line contains date pattern (likely a job entry)
                if re.search(r'\d{4}', line):
                    parts = line.split('-')
                    if len(parts) >= 2:
                        company = parts[0].strip()
                        if company and len(company) < 100:
                            experience.append({
                                "company": company,
                                "position": "Not Specified",
                                "duration": parts[-1].strip() if len(parts) > 1 else "Not Specified"
                            })
        
        return experience
    
    def parse_full_resume(self, file_path: str) -> Dict:
        """Complete resume parsing returning all extracted data"""
        # Parse file and get text
        text = self.parse_resume(file_path)
        
        if "Error" in text:
            return {
                "error": text,
                "basic_info": {
                    "name": "Not Found",
                    "email": "Not Found",
                    "phone": "Not Found"
                },
                "education": [],
                "experience": [],
                "skills": [],
                "pages": 0
            }
        
        # Extract all components
        return {
            "basic_info": {
                "name": self.extract_name(text),
                "email": self.extract_email(text),
                "phone": self.extract_phone(text)
            },
            "education": self.extract_education(text),
            "experience": self.extract_experience(text),
            "skills": self.extract_skills(text),
            "pages": self.get_page_count(file_path)
        }
